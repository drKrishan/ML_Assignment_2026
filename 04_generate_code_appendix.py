"""Generate a Word appendix that highlights key code segments for evaluation."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path(__file__).resolve().parent
APPENDIX_PATH = BASE_DIR / "BM173_Code_Appendix.docx"


def add_title(doc: Document, text: str) -> None:
    heading = doc.add_heading(text, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x4E)


def add_section_heading(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x4E)


def add_explanation(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)


def add_code_block(doc: Document, code: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.right_indent = Inches(0.2)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)


def build_appendix() -> None:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    add_title(doc, "BM173 Project Appendix: Key Code Segments")

    intro = (
        "This appendix highlights the most important implementation segments used in the "
        "Bank Marketing machine learning project. Each segment is accompanied by a short "
        "explanation so evaluators can quickly understand the logic, design decisions, and "
        "technical workflow."
    )
    add_explanation(doc, intro)

    # 1. Dataset loading and validation
    add_section_heading(doc, "1. Dataset Loading and Validation")
    add_explanation(
        doc,
        "Why it matters: This ensures the project always uses the correct real dataset and "
        "validates that the target column ('y') exists before modeling starts.",
    )
    add_code_block(
        doc,
        """from pathlib import Path
import pandas as pd

DATASET_PATH = Path('data') / 'bank-additional-full.csv'

if not DATASET_PATH.exists():
    raise FileNotFoundError(f'Dataset not found at {DATASET_PATH}')

df = pd.read_csv(DATASET_PATH, sep=';')
if 'y' not in df.columns:
    raise ValueError("Expected target column 'y' was not found in the dataset.")

print(df.shape)
print(df.head())""",
    )

    # 2. Target encoding and train-test split
    add_section_heading(doc, "2. Target Encoding and Stratified Train-Test Split")
    add_explanation(
        doc,
        "Why it matters: Converting the target to binary enables classification metrics. "
        "Stratified splitting preserves class imbalance proportions in both train and test sets.",
    )
    add_code_block(
        doc,
        """from sklearn.model_selection import train_test_split

df_model = df.copy()
df_model['y_binary'] = df_model['y'].map({'no': 0, 'yes': 1})

X = df_model.drop(columns=['y', 'y_binary'])
y = df_model['y_binary']

X_train, X_test, y_train, y_test = train_test_split(
    X, y,
    test_size=0.2,
    stratify=y,
    random_state=42
)""",
    )

    # 3. Preprocessing pipeline
    add_section_heading(doc, "3. ColumnTransformer Preprocessing Pipeline")
    add_explanation(
        doc,
        "Why it matters: Numeric and categorical features require different preprocessing. "
        "Using a single pipeline prevents data leakage and keeps transformations consistent "
        "during training and inference.",
    )
    add_code_block(
        doc,
        """from sklearn.compose import ColumnTransformer
from sklearn.impute import SimpleImputer
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler
import numpy as np

categorical_cols = X.select_dtypes(include=['object', 'string']).columns.tolist()
numeric_cols = X.select_dtypes(include=[np.number]).columns.tolist()

numeric_pipeline = Pipeline(steps=[
    ('imputer', SimpleImputer(strategy='median')),
    ('scaler', StandardScaler())
])

categorical_pipeline = Pipeline(steps=[
    ('imputer', SimpleImputer(strategy='most_frequent')),
    ('encoder', OneHotEncoder(handle_unknown='ignore', sparse_output=False))
])

preprocessor = ColumnTransformer(transformers=[
    ('numeric', numeric_pipeline, numeric_cols),
    ('categorical', categorical_pipeline, categorical_cols)
])""",
    )

    # 4. Model definitions (including ensembles)
    add_section_heading(doc, "4. Multi-Model Design (9 Models including Ensembles)")
    add_explanation(
        doc,
        "Why it matters: The project compares diverse algorithms to avoid model bias. "
        "It includes traditional models, neural network, and two ensemble strategies "
        "(Voting and Stacking) for performance improvement.",
    )
    add_code_block(
        doc,
        """models = {
    'Logistic Regression': Pipeline([
        ('preprocessor', preprocessor),
        ('model', LogisticRegression(max_iter=2000, class_weight='balanced', random_state=42))
    ]),
    'Random Forest': Pipeline([
        ('preprocessor', preprocessor),
        ('model', RandomForestClassifier(
            n_estimators=250,
            min_samples_leaf=5,
            class_weight='balanced_subsample',
            random_state=42,
            n_jobs=-1
        ))
    ]),
    'Neural Network (MLP)': Pipeline([
        ('preprocessor', preprocessor),
        ('model', MLPClassifier(
            hidden_layer_sizes=(128, 64, 32),
            activation='relu',
            solver='adam',
            alpha=0.001,
            learning_rate='adaptive',
            max_iter=300,
            early_stopping=True,
            validation_fraction=0.1,
            random_state=42
        ))
    ]),
    'Voting Ensemble': Pipeline([
        ('preprocessor', preprocessor),
        ('model', VotingClassifier(estimators=[...], voting='soft', n_jobs=-1))
    ]),
    'Stacking Ensemble': Pipeline([
        ('preprocessor', preprocessor),
        ('model', StackingClassifier(estimators=[...], final_estimator=LogisticRegression(...), cv=3, n_jobs=-1))
    ])
}""",
    )

    # 5. Training loop and metrics
    add_section_heading(doc, "5. Unified Training Loop and Evaluation Metrics")
    add_explanation(
        doc,
        "Why it matters: A single training loop applies the same evaluation protocol to every "
        "model, enabling fair comparison across Accuracy, Precision, Recall, F1-score, and ROC-AUC.",
    )
    add_code_block(
        doc,
        """from sklearn.metrics import (
    accuracy_score, precision_score, recall_score,
    f1_score, roc_auc_score, confusion_matrix, classification_report
)

evaluation_rows = []
detailed_results = {}

for model_name, pipeline in models.items():
    pipeline.fit(X_train, y_train)
    y_pred = pipeline.predict(X_test)
    y_prob = pipeline.predict_proba(X_test)[:, 1]

    metrics = {
        'Model': model_name,
        'Accuracy': accuracy_score(y_test, y_pred),
        'Precision': precision_score(y_test, y_pred, zero_division=0),
        'Recall': recall_score(y_test, y_pred, zero_division=0),
        'F1-Score': f1_score(y_test, y_pred, zero_division=0),
        'ROC-AUC': roc_auc_score(y_test, y_prob),
    }
    evaluation_rows.append(metrics)

    detailed_results[model_name] = {
        'confusion_matrix': confusion_matrix(y_test, y_pred),
        'classification_report': classification_report(y_test, y_pred, zero_division=0),
        'probabilities': y_prob,
    }

results_df = pd.DataFrame(evaluation_rows).set_index('Model').sort_values(
    by='F1-Score', ascending=False
)""",
    )

    # 6. ROC comparison plot
    add_section_heading(doc, "6. ROC Curve Comparison")
    add_explanation(
        doc,
        "Why it matters: ROC curves visualize trade-offs between true positive rate and false "
        "positive rate. This helps compare discrimination performance across all models.",
    )
    add_code_block(
        doc,
        """from sklearn.metrics import roc_curve
import matplotlib.pyplot as plt

fig, ax = plt.subplots(figsize=(11, 8))
for model_name in results_df.index:
    y_prob = detailed_results[model_name]['probabilities']
    fpr, tpr, _ = roc_curve(y_test, y_prob)
    ax.plot(fpr, tpr, label=f"{model_name} (AUC={results_df.loc[model_name, 'ROC-AUC']:.3f})")

ax.plot([0, 1], [0, 1], linestyle='--', color='black')
ax.set_title('ROC Curve Comparison')
ax.set_xlabel('False Positive Rate')
ax.set_ylabel('True Positive Rate')
ax.legend(loc='lower right')
plt.show()""",
    )

    # 7. Feature importance
    add_section_heading(doc, "7. Feature Importance for Interpretability")
    add_explanation(
        doc,
        "Why it matters: Feature importance identifies which transformed predictors most "
        "influence model decisions, strengthening business interpretation.",
    )
    add_code_block(
        doc,
        """pipeline = detailed_results['Random Forest']['pipeline']
preprocessor = pipeline.named_steps['preprocessor']
estimator = pipeline.named_steps['model']

feature_names = preprocessor.get_feature_names_out()
importance_df = pd.DataFrame({
    'Feature': feature_names,
    'Importance': estimator.feature_importances_
}).sort_values(by='Importance', ascending=False)

print(importance_df.head(15))""",
    )

    # 8. Single client prediction
    add_section_heading(doc, "8. New Client Prediction Across Multiple Models")
    add_explanation(
        doc,
        "Why it matters: This demonstrates practical deployment behavior by showing how each "
        "model responds to the same client profile.",
    )
    add_code_block(
        doc,
        """new_client = pd.DataFrame([{
    'age': 38,
    'job': 'admin.',
    'marital': 'married',
    'education': 'university.degree',
    'default': 'no',
    'housing': 'yes',
    'loan': 'no',
    'contact': 'cellular',
    'month': 'may',
    'day_of_week': 'thu',
    'duration': 220,
    'campaign': 2,
    'pdays': 999,
    'previous': 0,
    'poutcome': 'nonexistent',
    'emp.var.rate': 1.1,
    'cons.price.idx': 93.994,
    'cons.conf.idx': -36.4,
    'euribor3m': 4.857,
    'nr.employed': 5191.0,
}])

for model_name, pipeline in models.items():
    prob_yes = pipeline.predict_proba(new_client)[0, 1]
    print(f"{model_name}: P(yes)={prob_yes:.4f}")""",
    )

    # 9. File map for evaluators
    add_section_heading(doc, "9. File-to-Function Mapping")
    add_explanation(
        doc,
        "The implementation is organized by purpose so evaluators can trace logic quickly:",
    )
    add_code_block(
        doc,
        """01_download_dataset.py            -> Dataset download/verification
02_bank_marketing_analysis.py    -> Full ML workflow + metrics + visualizations
03_generate_report.py            -> Professional academic report generation
app.py                           -> Interactive dashboard for live predictions
Bank_Marketing_Project_Notebook.ipynb -> Unified educational workflow""",
    )

    conclusion = (
        "This appendix focuses on the critical implementation logic used for the analysis. "
        "It is intended to be attached as supplementary material so evaluators can verify "
        "technical depth, reproducibility, and model comparison methodology."
    )
    add_explanation(doc, conclusion)

    doc.save(APPENDIX_PATH)
    print(f"Appendix generated: {APPENDIX_PATH}")


if __name__ == "__main__":
    build_appendix()
