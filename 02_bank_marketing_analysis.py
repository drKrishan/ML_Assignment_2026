"""Bank Marketing classification analysis and Word report generation."""

from pathlib import Path
import textwrap

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from sklearn.compose import ColumnTransformer
from sklearn.ensemble import (
    GradientBoostingClassifier,
    RandomForestClassifier,
    StackingClassifier,
    VotingClassifier,
)
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import (
    accuracy_score,
    classification_report,
    confusion_matrix,
    f1_score,
    precision_score,
    recall_score,
    roc_auc_score,
    roc_curve,
)
from sklearn.model_selection import train_test_split
from sklearn.neighbors import KNeighborsClassifier
from sklearn.neural_network import MLPClassifier
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler
from sklearn.svm import SVC
from sklearn.tree import DecisionTreeClassifier

plt.style.use("seaborn-v0_8-darkgrid")
sns.set_palette("Set2")

BASE_DIR = Path(__file__).resolve().parent
DATASET_PATH = BASE_DIR / "data" / "bank-additional-full.csv"
VISUALS_DIR = BASE_DIR / "visualizations"
REPORT_PATH = BASE_DIR / "Bank_Marketing_Classification_Report.docx"
RESULTS_CSV_PATH = VISUALS_DIR / "model_results_summary.csv"
REPORT_TEXT_PATH = VISUALS_DIR / "classification_reports.txt"


def print_header(title: str) -> None:
    print("\n" + "=" * 90)
    print(title)
    print("=" * 90)


def load_dataset() -> pd.DataFrame:
    print_header("STEP 1: LOADING AND VALIDATING THE DATASET")
    if not DATASET_PATH.exists():
        raise FileNotFoundError(f"Dataset not found at {DATASET_PATH}")

    dataframe = pd.read_csv(DATASET_PATH, sep=";")
    print(f"Dataset loaded from: {DATASET_PATH}")
    print(f"Shape: {dataframe.shape[0]} rows x {dataframe.shape[1]} columns")
    print("Columns:")
    print(list(dataframe.columns))

    if "y" not in dataframe.columns:
        raise ValueError("Expected target column 'y' was not found in the dataset.")

    print("\nSample rows:")
    print(dataframe.head())
    print("\nData types:")
    print(dataframe.dtypes)
    print("\nMissing values per column:")
    print(dataframe.isna().sum())
    return dataframe


def create_eda_visualizations(dataframe: pd.DataFrame) -> None:
    print_header("STEP 2: EXPLORATORY DATA ANALYSIS")
    VISUALS_DIR.mkdir(exist_ok=True)

    target_counts = dataframe["y"].value_counts().sort_index()
    fig, axes = plt.subplots(1, 2, figsize=(14, 5))
    target_counts.plot(kind="bar", ax=axes[0], color=["#D95F02", "#1B9E77"])
    axes[0].set_title("Target Class Distribution")
    axes[0].set_xlabel("Subscription Outcome")
    axes[0].set_ylabel("Count")
    axes[0].tick_params(axis="x", rotation=0)

    target_counts.plot(
        kind="pie",
        ax=axes[1],
        autopct="%1.1f%%",
        colors=["#D95F02", "#1B9E77"],
        startangle=90,
    )
    axes[1].set_ylabel("")
    axes[1].set_title("Target Class Share")
    fig.tight_layout()
    fig.savefig(
        VISUALS_DIR / "01_target_distribution.png", dpi=300, bbox_inches="tight"
    )
    plt.close(fig)
    print("Saved 01_target_distribution.png")

    numeric_columns = dataframe.select_dtypes(include=[np.number]).columns.tolist()
    selected_numeric = numeric_columns[:6]
    fig, axes = plt.subplots(2, 3, figsize=(16, 9))
    axes = axes.ravel()
    for axis, column in zip(axes, selected_numeric):
        sns.histplot(dataframe[column], kde=True, ax=axis, color="#4C78A8")
        axis.set_title(f"Distribution of {column}")
    for axis in axes[len(selected_numeric) :]:
        axis.axis("off")
    fig.tight_layout()
    fig.savefig(
        VISUALS_DIR / "02_numeric_features_distribution.png",
        dpi=300,
        bbox_inches="tight",
    )
    plt.close(fig)
    print("Saved 02_numeric_features_distribution.png")

    categorical_columns = [
        column
        for column in dataframe.select_dtypes(include=["object", "string"]).columns
        if column != "y"
    ]
    selected_categorical = categorical_columns[:8]
    fig, axes = plt.subplots(2, 4, figsize=(18, 10))
    axes = axes.ravel()
    for axis, column in zip(axes, selected_categorical):
        order = dataframe[column].value_counts().index
        sns.countplot(data=dataframe, x=column, order=order, ax=axis, color="#72B7B2")
        axis.set_title(f"{column} distribution")
        axis.set_xlabel("")
        axis.tick_params(axis="x", rotation=45)
    for axis in axes[len(selected_categorical) :]:
        axis.axis("off")
    fig.tight_layout()
    fig.savefig(
        VISUALS_DIR / "03_categorical_features_distribution.png",
        dpi=300,
        bbox_inches="tight",
    )
    plt.close(fig)
    print("Saved 03_categorical_features_distribution.png")


def build_models(preprocessor: ColumnTransformer) -> dict[str, Pipeline]:
    return {
        "Logistic Regression": Pipeline(
            steps=[
                ("preprocessor", preprocessor),
                (
                    "model",
                    LogisticRegression(
                        max_iter=2000,
                        class_weight="balanced",
                        random_state=42,
                    ),
                ),
            ]
        ),
        "Decision Tree": Pipeline(
            steps=[
                ("preprocessor", preprocessor),
                (
                    "model",
                    DecisionTreeClassifier(
                        max_depth=8,
                        min_samples_leaf=20,
                        class_weight="balanced",
                        random_state=42,
                    ),
                ),
            ]
        ),
        "Random Forest": Pipeline(
            steps=[
                ("preprocessor", preprocessor),
                (
                    "model",
                    RandomForestClassifier(
                        n_estimators=250,
                        min_samples_leaf=5,
                        class_weight="balanced_subsample",
                        random_state=42,
                        n_jobs=-1,
                    ),
                ),
            ]
        ),
        "Support Vector Machine": Pipeline(
            steps=[
                ("preprocessor", preprocessor),
                (
                    "model",
                    SVC(
                        kernel="rbf",
                        probability=True,
                        class_weight="balanced",
                        random_state=42,
                    ),
                ),
            ]
        ),
        "K-Nearest Neighbors": Pipeline(
            steps=[
                ("preprocessor", preprocessor),
                ("model", KNeighborsClassifier(n_neighbors=15)),
            ]
        ),
        "Gradient Boosting": Pipeline(
            steps=[
                ("preprocessor", preprocessor),
                ("model", GradientBoostingClassifier(random_state=42)),
            ]
        ),
        "Neural Network (MLP)": Pipeline(
            steps=[
                ("preprocessor", preprocessor),
                (
                    "model",
                    MLPClassifier(
                        hidden_layer_sizes=(128, 64, 32),
                        activation="relu",
                        solver="adam",
                        alpha=0.001,
                        learning_rate="adaptive",
                        max_iter=300,
                        early_stopping=True,
                        validation_fraction=0.1,
                        random_state=42,
                    ),
                ),
            ]
        ),
        "Voting Ensemble": Pipeline(
            steps=[
                ("preprocessor", preprocessor),
                (
                    "model",
                    VotingClassifier(
                        estimators=[
                            (
                                "lr",
                                LogisticRegression(
                                    max_iter=2000,
                                    class_weight="balanced",
                                    random_state=42,
                                ),
                            ),
                            (
                                "rf",
                                RandomForestClassifier(
                                    n_estimators=250,
                                    min_samples_leaf=5,
                                    class_weight="balanced_subsample",
                                    random_state=42,
                                    n_jobs=-1,
                                ),
                            ),
                            ("gb", GradientBoostingClassifier(random_state=42)),
                            (
                                "mlp",
                                MLPClassifier(
                                    hidden_layer_sizes=(128, 64, 32),
                                    activation="relu",
                                    solver="adam",
                                    alpha=0.001,
                                    learning_rate="adaptive",
                                    max_iter=300,
                                    early_stopping=True,
                                    validation_fraction=0.1,
                                    random_state=42,
                                ),
                            ),
                        ],
                        voting="soft",
                        n_jobs=-1,
                    ),
                ),
            ]
        ),
        "Stacking Ensemble": Pipeline(
            steps=[
                ("preprocessor", preprocessor),
                (
                    "model",
                    StackingClassifier(
                        estimators=[
                            (
                                "lr",
                                LogisticRegression(
                                    max_iter=2000,
                                    class_weight="balanced",
                                    random_state=42,
                                ),
                            ),
                            (
                                "rf",
                                RandomForestClassifier(
                                    n_estimators=250,
                                    min_samples_leaf=5,
                                    class_weight="balanced_subsample",
                                    random_state=42,
                                    n_jobs=-1,
                                ),
                            ),
                            ("gb", GradientBoostingClassifier(random_state=42)),
                        ],
                        final_estimator=LogisticRegression(
                            max_iter=1000, class_weight="balanced", random_state=42
                        ),
                        cv=3,
                        n_jobs=-1,
                    ),
                ),
            ]
        ),
    }


def probability_scores(model: Pipeline, features: pd.DataFrame) -> np.ndarray:
    if hasattr(model, "predict_proba"):
        return model.predict_proba(features)[:, 1]
    if hasattr(model, "decision_function"):
        scores = model.decision_function(features)
        min_score = scores.min()
        max_score = scores.max()
        return (scores - min_score) / (max_score - min_score + 1e-9)
    raise AttributeError("Model does not support probability or decision scores.")


def evaluate_models(dataframe: pd.DataFrame) -> tuple[pd.DataFrame, dict, pd.Series]:
    print_header("STEP 3: DATA PREPARATION, TRAINING, AND EVALUATION")

    dataframe = dataframe.copy()
    dataframe["y_binary"] = dataframe["y"].map({"no": 0, "yes": 1})

    features = dataframe.drop(columns=["y", "y_binary"])
    target = dataframe["y_binary"]

    categorical_columns = features.select_dtypes(
        include=["object", "string"]
    ).columns.tolist()
    numeric_columns = features.select_dtypes(include=[np.number]).columns.tolist()

    print(f"Numeric columns: {len(numeric_columns)}")
    print(f"Categorical columns: {len(categorical_columns)}")
    print(
        f"Target distribution: {target.value_counts(normalize=True).round(4).to_dict()}"
    )

    X_train, X_test, y_train, y_test = train_test_split(
        features,
        target,
        test_size=0.2,
        stratify=target,
        random_state=42,
    )

    print(f"Training rows: {len(X_train)}")
    print(f"Testing rows: {len(X_test)}")

    numeric_pipeline = Pipeline(
        steps=[
            ("imputer", SimpleImputer(strategy="median")),
            ("scaler", StandardScaler()),
        ]
    )
    categorical_pipeline = Pipeline(
        steps=[
            ("imputer", SimpleImputer(strategy="most_frequent")),
            ("encoder", OneHotEncoder(handle_unknown="ignore", sparse_output=False)),
        ]
    )

    preprocessor = ColumnTransformer(
        transformers=[
            ("numeric", numeric_pipeline, numeric_columns),
            ("categorical", categorical_pipeline, categorical_columns),
        ]
    )

    models = build_models(preprocessor)
    evaluation_rows = []
    detailed_results = {}
    report_blocks = []

    for model_name, pipeline in models.items():
        print(f"Training {model_name}...")
        pipeline.fit(X_train, y_train)
        predictions = pipeline.predict(X_test)
        probabilities = probability_scores(pipeline, X_test)

        metrics = {
            "Model": model_name,
            "Accuracy": accuracy_score(y_test, predictions),
            "Precision": precision_score(y_test, predictions, zero_division=0),
            "Recall": recall_score(y_test, predictions, zero_division=0),
            "F1-Score": f1_score(y_test, predictions, zero_division=0),
            "ROC-AUC": roc_auc_score(y_test, probabilities),
        }
        evaluation_rows.append(metrics)
        detailed_results[model_name] = {
            "pipeline": pipeline,
            "predictions": predictions,
            "probabilities": probabilities,
            "confusion_matrix": confusion_matrix(y_test, predictions),
            "classification_report": classification_report(
                y_test,
                predictions,
                target_names=["No Subscription", "Subscription"],
                zero_division=0,
            ),
        }
        report_blocks.append(
            f"{model_name}\n{'-' * len(model_name)}\n{detailed_results[model_name]['classification_report']}\n"
        )
        print(
            f"  Accuracy={metrics['Accuracy']:.4f}, Precision={metrics['Precision']:.4f}, "
            f"Recall={metrics['Recall']:.4f}, F1={metrics['F1-Score']:.4f}, ROC-AUC={metrics['ROC-AUC']:.4f}"
        )

    results_df = pd.DataFrame(evaluation_rows).set_index("Model")
    results_df = results_df.astype(float).sort_values(by="F1-Score", ascending=False)
    RESULTS_CSV_PATH.parent.mkdir(exist_ok=True)
    results_df.to_csv(RESULTS_CSV_PATH)
    REPORT_TEXT_PATH.write_text("\n".join(report_blocks), encoding="utf-8")
    print(f"Saved {RESULTS_CSV_PATH.name}")
    print(f"Saved {REPORT_TEXT_PATH.name}")

    return results_df, detailed_results, y_test


def create_model_comparison_visuals(
    results_df: pd.DataFrame, detailed_results: dict, y_test: pd.Series
) -> None:
    print_header("STEP 4: CREATING MODEL COMPARISON VISUALS")

    metrics_to_plot = ["Accuracy", "Precision", "Recall", "F1-Score", "ROC-AUC"]
    heatmap_data = results_df[metrics_to_plot].copy().astype(float)

    fig, axes = plt.subplots(1, 2, figsize=(18, 7))
    heatmap_data.plot(kind="bar", ax=axes[0], width=0.85)
    axes[0].set_title("Performance Comparison Across Models")
    axes[0].set_xlabel("Model")
    axes[0].set_ylabel("Score")
    axes[0].tick_params(axis="x", rotation=35)
    axes[0].set_ylim(0, 1.05)
    axes[0].grid(axis="y", alpha=0.25)

    sns.heatmap(
        heatmap_data, annot=True, fmt=".3f", cmap="YlGnBu", ax=axes[1], vmin=0, vmax=1
    )
    axes[1].set_title("Metric Heatmap")
    fig.tight_layout()
    fig.savefig(VISUALS_DIR / "04_model_comparison.png", dpi=300, bbox_inches="tight")
    plt.close(fig)
    print("Saved 04_model_comparison.png")

    n_models = len(results_df)
    n_cols = 3
    n_rows = (n_models + n_cols - 1) // n_cols
    fig, axes = plt.subplots(n_rows, n_cols, figsize=(18, 6 * n_rows))
    axes = axes.ravel()
    for axis, model_name in zip(axes, results_df.index.tolist()):
        matrix = detailed_results[model_name]["confusion_matrix"]
        sns.heatmap(matrix, annot=True, fmt="d", cmap="Blues", cbar=False, ax=axis)
        axis.set_title(
            f"{model_name}\nF1 = {results_df.loc[model_name, 'F1-Score']:.3f}"
        )
        axis.set_xlabel("Predicted")
        axis.set_ylabel("Actual")
    # hide any unused axes
    for idx in range(n_models, len(axes)):
        axes[idx].set_visible(False)
    fig.tight_layout()
    fig.savefig(VISUALS_DIR / "05_confusion_matrices.png", dpi=300, bbox_inches="tight")
    plt.close(fig)
    print("Saved 05_confusion_matrices.png")

    fig, axis = plt.subplots(figsize=(11, 8))
    color_cycle = [
        "#1b9e77",
        "#d95f02",
        "#7570b3",
        "#e7298a",
        "#66a61e",
        "#e6ab02",
        "#a6761d",
        "#2166ac",
        "#b2182b",
    ]
    for color, model_name in zip(color_cycle, results_df.index.tolist()):
        probabilities = detailed_results[model_name]["probabilities"]
        fpr, tpr, _ = roc_curve(y_test, probabilities)
        axis.plot(
            fpr,
            tpr,
            label=f"{model_name} (AUC={results_df.loc[model_name, 'ROC-AUC']:.3f})",
            color=color,
            linewidth=2,
        )
    axis.plot([0, 1], [0, 1], linestyle="--", color="black", linewidth=1)
    axis.set_title("ROC Curve Comparison")
    axis.set_xlabel("False Positive Rate")
    axis.set_ylabel("True Positive Rate")
    axis.legend(loc="lower right")
    axis.grid(alpha=0.25)
    fig.tight_layout()
    fig.savefig(VISUALS_DIR / "06_roc_curves.png", dpi=300, bbox_inches="tight")
    plt.close(fig)
    print("Saved 06_roc_curves.png")


def create_feature_importance_visual(detailed_results: dict) -> pd.DataFrame:
    print_header("STEP 5: FEATURE IMPORTANCE ANALYSIS")

    top_tree_model = None
    for candidate in ["Random Forest", "Gradient Boosting", "Decision Tree"]:
        if candidate in detailed_results:
            top_tree_model = candidate
            break

    if top_tree_model is None:
        raise ValueError(
            "No tree-based model available for feature importance analysis."
        )

    pipeline = detailed_results[top_tree_model]["pipeline"]
    preprocessor = pipeline.named_steps["preprocessor"]
    estimator = pipeline.named_steps["model"]
    feature_names = preprocessor.get_feature_names_out()
    feature_importance = pd.DataFrame(
        {"Feature": feature_names, "Importance": estimator.feature_importances_}
    ).sort_values(by="Importance", ascending=False)

    top_features = feature_importance.head(15).copy()
    fig, axis = plt.subplots(figsize=(12, 8))
    sns.barplot(
        data=top_features, x="Importance", y="Feature", ax=axis, color="#4C78A8"
    )
    axis.set_title(f"Top 15 Feature Importances from {top_tree_model}")
    axis.set_xlabel("Importance")
    axis.set_ylabel("Feature")
    fig.tight_layout()
    fig.savefig(VISUALS_DIR / "07_feature_importance.png", dpi=300, bbox_inches="tight")
    plt.close(fig)
    print("Saved 07_feature_importance.png")

    feature_importance.to_csv(VISUALS_DIR / "feature_importance.csv", index=False)
    return feature_importance


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str) -> None:
    for block in textwrap.wrap(text, width=120, replace_whitespace=False):
        document.add_paragraph(block)


def add_image(document: Document, image_name: str, caption: str) -> None:
    image_path = VISUALS_DIR / image_name
    if image_path.exists():
        document.add_picture(str(image_path), width=Inches(6.5))
        caption_paragraph = document.add_paragraph(caption)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_word_report(
    dataframe: pd.DataFrame,
    results_df: pd.DataFrame,
    detailed_results: dict,
    feature_importance: pd.DataFrame,
) -> None:
    print_header("STEP 6: GENERATING THE WORD REPORT")

    best_model = results_df.index[0]
    best_metrics = results_df.iloc[0]
    class_balance = dataframe["y"].value_counts(normalize=True).mul(100).round(2)

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = document.add_heading(
        "Bank Marketing Classification Analysis Report", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(
        "BM 173 Applications of Machine Learning | Individual analytical report"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(document, "1. Introduction", level=1)
    add_paragraph(
        document,
        "This report evaluates a real business classification problem using the UCI Bank Marketing dataset. The business objective is to predict whether a client will subscribe to a term deposit after a direct marketing campaign. This problem is highly relevant to business analytics because accurate classification allows firms to target likely responders, reduce wasted campaign effort, and improve conversion efficiency.",
    )
    add_paragraph(
        document,
        f"The dataset contains {len(dataframe):,} customer records and {dataframe.shape[1] - 1} predictor variables. The target is imbalanced: {class_balance['no']}% of observations are non-subscribers and {class_balance['yes']}% are subscribers. This imbalance makes simple accuracy insufficient on its own and justifies the use of precision, recall, F1-score, and ROC-AUC for evaluation.",
    )

    add_heading(document, "2. Dataset and Business Context", level=1)
    add_paragraph(
        document,
        "The variables represent demographic information, financial status, campaign contact history, and macroeconomic indicators. Examples include age, job, marital status, education, loan status, contact channel, campaign duration, number of previous contacts, and market conditions such as euribor and employment indicators. Together these variables make the dataset suitable for a realistic business segmentation and response prediction exercise.",
    )
    add_image(
        document,
        "01_target_distribution.png",
        "Figure 1. Target distribution of term deposit subscriptions.",
    )

    add_heading(document, "3. Data Exploration and Preparation", level=1)
    add_paragraph(
        document,
        "Initial exploration confirmed that the dataset is structurally complete and contains no missing values in the supplied file. Numerical variables were reviewed through distribution plots, while categorical fields were examined using count-based charts. The data includes both numeric and text fields, so preprocessing was implemented with separate transformations for each type.",
    )
    add_paragraph(
        document,
        "A preprocessing pipeline was built with median imputation and standardisation for numeric features and most-frequent imputation plus one-hot encoding for categorical features. This approach preserves the business meaning of categorical attributes while producing a machine-learning-ready matrix for all models. The dataset was split into 80% training data and 20% test data using stratified sampling so that the class distribution remained stable across both subsets.",
    )
    add_image(
        document,
        "02_numeric_features_distribution.png",
        "Figure 2. Distributions of selected numeric variables.",
    )
    add_image(
        document,
        "03_categorical_features_distribution.png",
        "Figure 3. Distributions of selected categorical variables.",
    )

    add_heading(document, "4. Model Development", level=1)
    add_paragraph(
        document,
        "Nine classification algorithms were trained, spanning linear, distance-based, kernel-based, tree-based, deep learning, and combined ensemble strategies: Logistic Regression, Decision Tree, Random Forest, Support Vector Machine, K-Nearest Neighbors, Gradient Boosting, Neural Network (MLP), Voting Ensemble, and Stacking Ensemble. The MLP was configured with three hidden layers (128-64-32 neurons), ReLU activations, the Adam optimiser, and early stopping. The Voting Ensemble combines Logistic Regression, Random Forest, Gradient Boosting, and MLP using soft voting (averaged predicted probabilities), reducing variance by leveraging the complementary strengths of diverse model types. The Stacking Ensemble uses Logistic Regression, Random Forest, and Gradient Boosting as base learners with a Logistic Regression meta-learner trained via 3-fold cross-validation, learning the optimal way to weight each base model's predictions. Class balancing was introduced where appropriate to reduce bias toward the majority class.",
    )

    add_heading(document, "5. Model Evaluation", level=1)
    add_paragraph(
        document,
        "Model quality was assessed using accuracy, precision, recall, F1-score, and ROC-AUC. Accuracy measures overall correctness, precision measures how reliable positive subscription predictions are, recall measures how many true subscribers were captured, F1-score balances precision and recall, and ROC-AUC measures ranking quality across decision thresholds. For this business problem, F1-score and ROC-AUC are especially useful because the target distribution is imbalanced.",
    )

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Model", "Accuracy", "Precision", "Recall", "F1-Score", "ROC-AUC"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for model_name, row in results_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = model_name
        row_cells[1].text = f"{row['Accuracy']:.4f}"
        row_cells[2].text = f"{row['Precision']:.4f}"
        row_cells[3].text = f"{row['Recall']:.4f}"
        row_cells[4].text = f"{row['F1-Score']:.4f}"
        row_cells[5].text = f"{row['ROC-AUC']:.4f}"

    add_paragraph(
        document,
        f"The best-performing model by F1-score was {best_model}, with accuracy of {best_metrics['Accuracy']:.4f}, precision of {best_metrics['Precision']:.4f}, recall of {best_metrics['Recall']:.4f}, F1-score of {best_metrics['F1-Score']:.4f}, and ROC-AUC of {best_metrics['ROC-AUC']:.4f}. This indicates that the model offered the strongest balance between identifying likely subscribers and limiting false positives.",
    )
    add_image(
        document,
        "04_model_comparison.png",
        "Figure 4. Comparison of model performance across evaluation metrics.",
    )
    add_image(
        document,
        "05_confusion_matrices.png",
        "Figure 5. Confusion matrices for all classification models.",
    )
    add_image(
        document,
        "06_roc_curves.png",
        "Figure 6. ROC curve comparison across all models.",
    )

    add_heading(document, "6. Discussion of Results", level=1)
    add_paragraph(
        document,
        f"{best_model} produced the strongest overall balance on the test set. In business terms, this means the model is better suited than the alternatives for prioritising customers who are more likely to subscribe while maintaining reasonable discrimination between responders and non-responders. Random Forest and Gradient Boosting also performed strongly, which suggests that non-linear relationships and interaction effects matter in this marketing problem.",
    )
    add_paragraph(
        document,
        "The results also highlight an important managerial trade-off. A model with higher precision reduces the cost of unnecessary follow-ups, but a model with higher recall captures more potential subscribers. The preferred model depends on campaign economics. If outreach is expensive, higher precision may be prioritised. If losing a potential subscriber is more costly, higher recall may be preferred even at the expense of more false positives.",
    )

    add_heading(document, "7. Feature Importance and Business Interpretation", level=1)
    top_five = feature_importance.head(5)[["Feature", "Importance"]].values.tolist()
    top_five_text = ", ".join([f"{name} ({score:.4f})" for name, score in top_five])
    add_paragraph(
        document,
        f"Feature importance analysis shows that the most influential transformed predictors were: {top_five_text}. These results suggest that campaign interaction quality, prior contact outcomes, and customer profile variables all contribute materially to subscription decisions. In practice, this means a bank can improve campaign targeting by combining behavioural, demographic, and economic context rather than relying on a single factor.",
    )
    add_image(
        document,
        "07_feature_importance.png",
        "Figure 7. Top transformed features driving the best tree-based model.",
    )

    add_heading(document, "8. Limitations and Potential Improvements", level=1)
    add_paragraph(
        document,
        "Several limitations should be acknowledged. First, the target class is heavily imbalanced, so performance could be further improved with threshold tuning, resampling methods such as SMOTE, or cost-sensitive learning. Second, the dataset includes call duration, which is highly informative but may not always be known in advance if the goal is to predict campaign success before a call takes place. Third, this analysis used a single train-test split; future work could add cross-validation and hyperparameter optimisation for more robust model selection.",
    )
    add_paragraph(
        document,
        "Recommended next improvements include hyperparameter tuning for the top three models, threshold optimisation aligned to campaign cost assumptions, removal or separate treatment of post-contact variables such as duration, and a segmented analysis by customer profile to support more targeted marketing strategy.",
    )

    add_heading(document, "9. Conclusion", level=1)
    add_paragraph(
        document,
        f"This classification exercise demonstrates that machine learning can support direct marketing decisions in a real business context. Using {len(dataframe):,} observations from the Bank Marketing dataset, multiple models were compared using a consistent preprocessing pipeline and appropriate evaluation metrics. {best_model} achieved the best balance of predictive performance and is the strongest candidate for identifying high-potential customers for a term deposit campaign. The overall analysis shows clear business value: better prioritisation, more efficient campaign spending, and a stronger evidence base for customer targeting decisions.",
    )

    document.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(document, "Appendix: Classification Reports", level=1)
    for model_name in results_df.index.tolist():
        add_heading(document, model_name, level=2)
        document.add_paragraph(detailed_results[model_name]["classification_report"])

    document.save(REPORT_PATH)
    print(f"Word report saved to: {REPORT_PATH}")


def main() -> None:
    dataframe = load_dataset()
    create_eda_visualizations(dataframe)
    results_df, detailed_results, y_test = evaluate_models(dataframe)
    create_model_comparison_visuals(results_df, detailed_results, y_test)
    feature_importance = create_feature_importance_visual(detailed_results)
    build_word_report(dataframe, results_df, detailed_results, feature_importance)

    print_header("ANALYSIS COMPLETE")
    print(results_df.round(4))
    print(f"\nArtifacts saved in: {VISUALS_DIR}")
    print(f"Report saved at: {REPORT_PATH}")


if __name__ == "__main__":
    main()
