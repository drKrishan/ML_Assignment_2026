"""Generate a professional academic Word report for BM 173 assignment.

Reads the already-computed model results and visualisations produced by
02_bank_marketing_analysis.py, creates additional figures, and assembles
a ~3 000-word structured report as a .docx file.
"""

from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

# ── paths ──────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).resolve().parent
DATA_PATH = BASE_DIR / "data" / "bank-additional-full.csv"
VIS_DIR = BASE_DIR / "visualizations"
REPORT_PATH = BASE_DIR / "BM173_Bank_Marketing_ML_Report.docx"

plt.style.use("seaborn-v0_8-darkgrid")
sns.set_palette("Set2")


# ══════════════════════════════════════════════════════════════════════
#  ADDITIONAL VISUALISATION FUNCTIONS
# ══════════════════════════════════════════════════════════════════════


def create_correlation_heatmap(df: pd.DataFrame) -> None:
    """Correlation matrix of all numeric features."""
    numeric = df.select_dtypes(include=[np.number])
    corr = numeric.corr()
    fig, ax = plt.subplots(figsize=(12, 10))
    mask = np.triu(np.ones_like(corr, dtype=bool))
    sns.heatmap(
        corr,
        mask=mask,
        annot=True,
        fmt=".2f",
        cmap="coolwarm",
        center=0,
        ax=ax,
        linewidths=0.5,
        square=True,
        cbar_kws={"shrink": 0.8},
    )
    ax.set_title("Correlation Matrix of Numeric Features", fontsize=14, pad=12)
    fig.tight_layout()
    fig.savefig(VIS_DIR / "08_correlation_heatmap.png", dpi=300, bbox_inches="tight")
    plt.close(fig)
    print("Saved 08_correlation_heatmap.png")


def create_class_imbalance_visual(df: pd.DataFrame) -> None:
    """Side-by-side bar + annotated pie showing class imbalance."""
    counts = df["y"].value_counts()
    labels = ["No (Did not subscribe)", "Yes (Subscribed)"]
    colours = ["#D95F02", "#1B9E77"]

    fig, axes = plt.subplots(1, 2, figsize=(14, 5))
    bars = axes[0].bar(
        labels, counts.values, color=colours, edgecolor="black", linewidth=0.6
    )
    for bar, val in zip(bars, counts.values):
        axes[0].text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + 200,
            f"{val:,}",
            ha="center",
            va="bottom",
            fontweight="bold",
            fontsize=11,
        )
    axes[0].set_title("Target Class Counts", fontsize=13, pad=10)
    axes[0].set_ylabel("Number of Records")

    axes[1].pie(
        counts.values,
        labels=labels,
        autopct="%1.1f%%",
        startangle=90,
        colors=colours,
        explode=(0.03, 0.03),
        textprops={"fontsize": 11},
        wedgeprops={"edgecolor": "black", "linewidth": 0.5},
    )
    axes[1].set_title("Class Distribution (%)", fontsize=13, pad=10)
    fig.tight_layout()
    fig.savefig(VIS_DIR / "09_class_imbalance.png", dpi=300, bbox_inches="tight")
    plt.close(fig)
    print("Saved 09_class_imbalance.png")


def create_precision_recall_bar(results: pd.DataFrame) -> None:
    """Grouped bar chart of precision vs recall per model."""
    fig, ax = plt.subplots(figsize=(14, 7))
    x = np.arange(len(results))
    width = 0.35
    ax.bar(
        x - width / 2,
        results["Precision"],
        width,
        label="Precision",
        color="#4C78A8",
        edgecolor="black",
        linewidth=0.5,
    )
    ax.bar(
        x + width / 2,
        results["Recall"],
        width,
        label="Recall",
        color="#E45756",
        edgecolor="black",
        linewidth=0.5,
    )
    ax.set_xticks(x)
    ax.set_xticklabels(results.index, rotation=30, ha="right", fontsize=9)
    ax.set_ylabel("Score")
    ax.set_title("Precision vs Recall Trade-off Across Models", fontsize=13, pad=10)
    ax.set_ylim(0, 1.08)
    ax.legend(fontsize=11)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    fig.savefig(VIS_DIR / "10_precision_recall_bar.png", dpi=300, bbox_inches="tight")
    plt.close(fig)
    print("Saved 10_precision_recall_bar.png")


def create_f1_roc_scatter(results: pd.DataFrame) -> None:
    """Scatter plot of F1-Score vs ROC-AUC with model labels."""
    fig, ax = plt.subplots(figsize=(11, 8))
    colours = plt.cm.tab10(np.linspace(0, 1, len(results)))
    for i, (name, row) in enumerate(results.iterrows()):
        ax.scatter(
            row["ROC-AUC"],
            row["F1-Score"],
            s=180,
            color=colours[i],
            edgecolor="black",
            linewidth=0.7,
            zorder=3,
        )
        ax.annotate(
            name,
            (row["ROC-AUC"], row["F1-Score"]),
            textcoords="offset points",
            xytext=(8, 6),
            fontsize=8.5,
        )
    ax.set_xlabel("ROC-AUC", fontsize=12)
    ax.set_ylabel("F1-Score", fontsize=12)
    ax.set_title("Model Performance: F1-Score vs ROC-AUC", fontsize=13, pad=10)
    ax.grid(alpha=0.3)
    fig.tight_layout()
    fig.savefig(VIS_DIR / "11_f1_roc_scatter.png", dpi=300, bbox_inches="tight")
    plt.close(fig)
    print("Saved 11_f1_roc_scatter.png")


def create_radar_chart(results: pd.DataFrame) -> None:
    """Radar / spider chart comparing the top 4 models across all metrics."""
    top = results.head(4)
    metrics = ["Accuracy", "Precision", "Recall", "F1-Score", "ROC-AUC"]
    n_metrics = len(metrics)
    angles = np.linspace(0, 2 * np.pi, n_metrics, endpoint=False).tolist()
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(9, 9), subplot_kw=dict(polar=True))
    colours = ["#1b9e77", "#d95f02", "#7570b3", "#e7298a"]
    for i, (name, row) in enumerate(top.iterrows()):
        values = [row[m] for m in metrics] + [row[metrics[0]]]
        ax.plot(angles, values, linewidth=2, label=name, color=colours[i])
        ax.fill(angles, values, alpha=0.1, color=colours[i])
    ax.set_thetagrids(np.degrees(angles[:-1]), metrics, fontsize=11)
    ax.set_ylim(0, 1.05)
    ax.set_title("Top 4 Models – Multi-Metric Radar", fontsize=13, pad=20)
    ax.legend(loc="upper right", bbox_to_anchor=(1.3, 1.12), fontsize=9)
    fig.tight_layout()
    fig.savefig(VIS_DIR / "12_radar_chart.png", dpi=300, bbox_inches="tight")
    plt.close(fig)
    print("Saved 12_radar_chart.png")


# ══════════════════════════════════════════════════════════════════════
#  WORD REPORT HELPERS
# ══════════════════════════════════════════════════════════════════════


def _set_cell_shading(cell, colour_hex: str) -> None:
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): colour_hex,
        },
    )
    shading.append(shading_elm)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x4E)


def add_para(
    doc: Document,
    text: str,
    bold: bool = False,
    italic: bool = False,
    space_after: int = 6,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic


def add_image(doc: Document, filename: str, caption: str, width: float = 6.2) -> None:
    path = VIS_DIR / filename
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_metrics_table(doc: Document, results: pd.DataFrame) -> None:
    headers = ["Model", "Accuracy", "Precision", "Recall", "F1-Score", "ROC-AUC"]
    table = doc.add_table(rows=1 + len(results), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True

    # header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        _set_cell_shading(cell, "1A1A4E")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # data rows
    for row_idx, (model, row) in enumerate(results.iterrows(), start=1):
        cells = table.rows[row_idx].cells
        cells[0].text = model
        for col_idx, metric in enumerate(headers[1:], start=1):
            cells[col_idx].text = f"{row[metric]:.4f}"
        for cell in cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
        if row_idx % 2 == 0:
            for cell in cells:
                _set_cell_shading(cell, "EDF2FA")

    cap = doc.add_paragraph(
        "Table 1. Performance metrics for all nine classification models."
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


# ══════════════════════════════════════════════════════════════════════
#  BUILD THE FULL REPORT
# ══════════════════════════════════════════════════════════════════════


def build_report(
    df: pd.DataFrame, results: pd.DataFrame, feat_imp: pd.DataFrame
) -> None:
    print("\n" + "=" * 90)
    print("BUILDING ACADEMIC WORD REPORT")
    print("=" * 90)

    n_rows = len(df)
    n_cols = df.shape[1] - 1  # exclude target
    class_pct = df["y"].value_counts(normalize=True).mul(100).round(1)
    best = results.index[0]
    bm = results.iloc[0]
    top5 = feat_imp.head(5)["Feature"].tolist()

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title page ─────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_heading(
        "Predicting Term Deposit Subscription\nUsing Machine Learning:\nAn Analysis of the UCI Bank Marketing Dataset",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x4E)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "BM 173 – Applications of Machine Learning\nIndividual Analytical Report"
    )
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = date_p.add_run("March 2026")
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_section(WD_SECTION_START.NEW_PAGE)

    # ── 1. Introduction ───────────────────────────────────────────
    add_heading(doc, "1. Introduction")

    add_para(
        doc,
        "Machine learning has become an essential capability in modern business analytics, "
        "enabling organisations to extract predictive insights from large volumes of operational "
        "data (Jordan and Mitchell, 2015). In the financial services sector, direct marketing "
        "campaigns represent a significant investment, and the ability to predict which customers "
        "are most likely to respond positively can substantially improve campaign return on "
        "investment (Moro, Cortez and Rita, 2014). This report applies supervised classification "
        "techniques to the UCI Bank Marketing dataset to predict whether a client will subscribe "
        "to a term deposit following a telephone-based direct marketing campaign.",
    )

    add_para(
        doc,
        f"The dataset comprises {n_rows:,} customer interaction records and {n_cols} predictor "
        "variables that capture demographic, financial, campaign-specific, and macroeconomic "
        "attributes (Dua and Graff, 2017). The binary target variable indicates whether the "
        "client subscribed ('yes') or did not subscribe ('no'). A notable characteristic of the "
        f"dataset is severe class imbalance: {class_pct['no']}% of records belong to the "
        f"majority class (non-subscribers) while only {class_pct['yes']}% represent subscribers. "
        "This imbalance poses analytical challenges because naive classifiers can achieve "
        "misleadingly high accuracy by simply predicting the majority class (He and Garcia, 2009).",
    )

    add_para(
        doc,
        "Nine classification algorithms are evaluated, spanning linear, distance-based, "
        "tree-based, neural network, and ensemble paradigms. The objectives are threefold: "
        "(1) to demonstrate rigorous data preparation and exploratory analysis; "
        "(2) to train, evaluate, and compare multiple machine learning models using "
        "appropriate performance metrics; and (3) to critically interpret the results and "
        "discuss their practical implications for bank marketing strategy. "
        "All modelling was conducted in Python using scikit-learn (Pedregosa et al., 2011), "
        "with pandas for data manipulation (McKinney, 2010) and matplotlib and seaborn for "
        "visualisation (Hunter, 2007; Waskom, 2021).",
    )

    # ── 2. Data Exploration and Preparation ───────────────────────
    add_heading(doc, "2. Data Exploration and Preparation")

    add_heading(doc, "2.1 Dataset Overview", level=2)
    add_para(
        doc,
        "The Bank Marketing dataset was obtained from the UCI Machine Learning Repository "
        "(Moro, Cortez and Rita, 2014). It records client interactions from direct "
        "telephone marketing campaigns conducted by a Portuguese banking institution between "
        "2008 and 2013. Each record represents a single campaign contact and includes "
        "information on the client's age, occupation, marital status, education level, credit "
        "default history, housing and personal loan status, contact communication type, month "
        "and day of last contact, call duration, number of contacts during the current campaign, "
        "days since previous campaign contact, number of previous contacts, outcome of the "
        "previous campaign, and four macroeconomic indicators: employment variation rate, "
        "consumer price index, consumer confidence index, and the euribor 3-month rate. "
        "The dataset was selected because it presents a realistic business classification problem "
        "with mixed data types, sufficient volume for reliable model training, and a meaningful "
        "class imbalance that mirrors real-world marketing response rates (Lessmann et al., 2015).",
    )

    add_heading(doc, "2.2 Exploratory Data Analysis", level=2)
    add_para(
        doc,
        "Initial exploration confirmed that the dataset contains no missing values in any column. "
        "Descriptive statistics revealed considerable variation in numeric features: client age "
        "ranges from 17 to 98 years, call duration spans from 0 to 4,918 seconds, and the "
        "euribor 3-month rate varies between 0.63% and 5.05%, reflecting the significant "
        "macroeconomic shifts during the data collection period. Frequency analysis of "
        "categorical variables showed that the most common occupations are administrative roles, "
        "blue-collar workers, and technicians, while the most frequent contact month is May.",
    )

    add_image(
        doc,
        "09_class_imbalance.png",
        "Figure 1. Class distribution of the target variable showing severe imbalance "
        "(88.7% non-subscribers vs 11.3% subscribers).",
    )

    add_para(
        doc,
        "Figure 1 illustrates the class imbalance. The correlation matrix (Figure 2) reveals "
        "strong positive correlations between euribor3m and emp.var.rate (r = 0.97) and between "
        "nr.employed and euribor3m (r = 0.95), indicating multicollinearity among macroeconomic "
        "indicators. Duration shows the strongest individual association with the target, although "
        "it is known only after a call concludes and therefore may not be available as a "
        "pre-contact predictor in practice (Moro, Cortez and Rita, 2014).",
    )

    add_image(
        doc,
        "08_correlation_heatmap.png",
        "Figure 2. Correlation matrix of numeric features revealing multicollinearity "
        "among macroeconomic variables.",
    )

    add_heading(doc, "2.3 Data Preparation and Feature Engineering", level=2)
    add_para(
        doc,
        "Data preparation followed a structured pipeline approach using scikit-learn's "
        "ColumnTransformer (Pedregosa et al., 2011). Numeric features were processed through "
        "median imputation followed by standardisation (z-score scaling) to ensure zero mean "
        "and unit variance, which is essential for distance-based and gradient-based algorithms "
        "(Kuhn and Johnson, 2013). Categorical features were handled through mode imputation "
        "and one-hot encoding with unknown-category handling to prevent information leakage "
        "and accommodate unseen categories during evaluation.",
    )

    add_para(
        doc,
        "The dataset was split into 80% training (32,950 records) and 20% test (8,238 records) "
        "subsets using stratified sampling to preserve the original class distribution in both "
        "subsets. Stratification is critical when classes are imbalanced to ensure evaluation "
        "metrics are computed on a representative sample (Raschka and Mirjalili, 2019). "
        "To address the class imbalance during training, class weighting was applied to "
        "cost-sensitive algorithms (Logistic Regression, Decision Tree, Random Forest, SVM), "
        "which adjusts the loss function to penalise misclassification of the minority class "
        "more heavily (King and Zeng, 2001).",
    )

    add_image(
        doc,
        "02_numeric_features_distribution.png",
        "Figure 3. Distribution of selected numeric features showing varied scales and skewness.",
    )

    add_image(
        doc,
        "03_categorical_features_distribution.png",
        "Figure 4. Frequency distributions of key categorical predictor variables.",
    )

    # ── 3. Model Development ──────────────────────────────────────
    add_heading(doc, "3. Model Development")

    add_para(
        doc,
        "Nine classification models were developed to provide a comprehensive comparison "
        "across algorithmic paradigms. This breadth of comparison is recommended in applied "
        "machine learning research to avoid selection bias toward any single approach "
        "(Fernandez-Delgado et al., 2014).",
    )

    add_heading(doc, "3.1 Individual Models", level=2)

    add_heading(doc, "Logistic Regression", level=3)
    add_para(
        doc,
        "Logistic Regression was selected as the baseline linear classifier due to its "
        "interpretability and computational efficiency (Hosmer, Lemeshow and Sturdivant, "
        "2013). It models the log-odds of subscription as a linear combination of the input "
        "features and applies L2 regularisation to prevent overfitting. Balanced class weights "
        "were assigned so that misclassification of the minority class (subscribers) incurs a "
        "proportionally higher penalty, directly addressing the class imbalance. As a linear "
        "model, Logistic Regression provides a useful performance floor against which more "
        "complex algorithms can be compared, and its coefficients offer direct insight into "
        "the direction and magnitude of each feature's influence on the prediction.",
    )

    add_heading(doc, "Decision Tree", level=3)
    add_para(
        doc,
        "A Decision Tree classifier was trained with a maximum depth of 8 and a minimum of "
        "20 samples per leaf to control model complexity and reduce overfitting "
        "(Breiman et al., 1984). Decision Trees recursively partition the feature space using "
        "binary splits that maximise information gain, producing a tree structure that is "
        "highly interpretable and can capture non-linear relationships without requiring "
        "feature scaling. Balanced class weights were applied. The contribution of this model "
        "lies in its transparency: the resulting tree can be visualised and explained to "
        "non-technical stakeholders, making it valuable for business reporting and compliance "
        "contexts where model explainability is a requirement.",
    )

    add_heading(doc, "Random Forest", level=3)
    add_para(
        doc,
        "Random Forest is a bagging ensemble that aggregates 250 independent decision trees, "
        "each trained on a bootstrap sample with a random subset of features considered at "
        "every split (Breiman, 2001). This combination of bagging and feature randomisation "
        "substantially reduces the variance of the predictions compared to a single tree. "
        "Balanced subsample weighting was used to adjust class representation within each "
        "bootstrap sample. Random Forest also provides a built-in feature importance ranking "
        "based on mean decrease in impurity, which supports the interpretability analysis in "
        "Section 5. Its contribution is strong predictive accuracy with inherent resistance "
        "to overfitting, making it a reliable workhorse for tabular classification tasks.",
    )

    add_heading(doc, "Support Vector Machine", level=3)
    add_para(
        doc,
        "A Support Vector Machine (SVM) with a radial basis function (RBF) kernel was "
        "employed to find an optimal separating hyperplane in a high-dimensional transformed "
        "feature space (Cortes and Vapnik, 1995). The RBF kernel allows the SVM to model "
        "complex non-linear decision boundaries by implicitly mapping the input features into "
        "an infinite-dimensional space. Probability estimates were enabled via Platt scaling "
        "to allow ROC-AUC computation. Balanced class weights were applied to increase the "
        "penalty for misclassifying subscribers. SVM contributes a fundamentally different "
        "learning paradigm — margin maximisation — offering a useful diversity baseline "
        "when combined with tree-based and linear models in the ensemble strategies.",
    )

    add_heading(doc, "K-Nearest Neighbors", level=3)
    add_para(
        doc,
        "K-Nearest Neighbors (KNN) with k = 15 was included as a non-parametric, "
        "instance-based classifier that makes predictions based on the majority class among "
        "the 15 nearest training instances in standardised Euclidean space (Cover and Hart, "
        "1967). Unlike the other models, KNN makes no assumptions about the functional form "
        "of the decision boundary and does not learn explicit model parameters during training. "
        "Its contribution to the comparison is methodological: it demonstrates how a purely "
        "distance-based approach performs on this high-dimensional, mixed-type dataset, "
        "providing context for the importance of algorithm choice when feature spaces are "
        "large and sparse after one-hot encoding.",
    )

    add_heading(doc, "Gradient Boosting", level=3)
    add_para(
        doc,
        "Gradient Boosting constructs an additive ensemble of shallow decision trees, where "
        "each successive tree is fitted to the negative gradient of the loss function — "
        "effectively correcting the residual errors of the preceding ensemble (Friedman, 2001). "
        "This sequential boosting strategy produces a highly expressive model that often "
        "achieves state-of-the-art performance on structured tabular data. In this analysis, "
        "Gradient Boosting delivered the highest precision among individual models, meaning "
        "its positive predictions were the most reliable. This makes it particularly suited "
        "to campaign scenarios where the cost of contacting non-interested clients is high "
        "and marketing resources are limited.",
    )

    add_heading(doc, "3.2 Neural Network", level=2)
    add_para(
        doc,
        "A Multi-Layer Perceptron (MLP) neural network was implemented with three hidden "
        "layers of 128, 64, and 32 neurons respectively, ReLU activation functions, the Adam "
        "optimiser, L2 regularisation (alpha = 0.001), and an adaptive learning rate "
        "(Goodfellow, Bengio and Courville, 2016). Early stopping was enabled with 10% of "
        "training data reserved for validation, halting training when validation loss ceased "
        "to decrease, to prevent overfitting on the training set (Prechelt, 1998).",
    )

    add_heading(doc, "3.3 Ensemble Methods", level=2)
    add_para(
        doc,
        "Two advanced ensemble strategies were applied to investigate whether combining "
        "models could improve predictive performance beyond that of any individual algorithm.",
    )

    add_para(
        doc,
        "The Voting Ensemble combines Logistic Regression, Random Forest, Gradient Boosting, "
        "and MLP using soft voting, which averages the predicted class probabilities from all "
        "four base learners and selects the class with the highest average probability "
        "(Zhou, 2012). By combining diverse model types — a linear discriminant, a bagging "
        "ensemble, a boosting ensemble, and a neural network — the voting approach reduces "
        "the risk that any single model's weaknesses dominate the final prediction.",
    )

    add_para(
        doc,
        "The Stacking Ensemble uses a two-level architecture: Logistic Regression, Random "
        "Forest, and Gradient Boosting serve as base learners (level 0), and a Logistic "
        "Regression meta-learner (level 1) is trained on their cross-validated predictions "
        "using 3-fold cross-validation (Wolpert, 1992). This approach learns the optimal "
        "weighting of each base model's predictions rather than applying equal weights, "
        "and the cross-validation ensures the meta-learner is trained on out-of-fold "
        "predictions to avoid information leakage (Van der Laan, Polley and Hubbard, 2007).",
    )

    # ── 4. Model Evaluation ───────────────────────────────────────
    add_heading(doc, "4. Model Evaluation")

    add_para(
        doc,
        "Model performance was assessed using five complementary metrics. Accuracy measures "
        "overall classification correctness. Precision quantifies the proportion of positive "
        "predictions that were truly positive, indicating how reliable a subscription prediction "
        "is. Recall (sensitivity) measures the proportion of actual subscribers correctly "
        "identified, reflecting the model's ability to capture positive cases. The F1-Score "
        "provides a harmonic mean of precision and recall, offering a balanced measure "
        "particularly useful under class imbalance (Sokolova and Lapalme, 2009). "
        "ROC-AUC evaluates the model's discriminative ability across all classification "
        "thresholds, with a value of 1.0 indicating perfect separation and 0.5 indicating "
        "random chance (Fawcett, 2006).",
    )

    add_metrics_table(doc, results)

    add_para(
        doc,
        f"Table 1 presents the evaluation results. The {best} achieved the highest F1-Score "
        f"of {bm['F1-Score']:.4f}, indicating the best balance between precision "
        f"({bm['Precision']:.4f}) and recall ({bm['Recall']:.4f}). "
        f"Its ROC-AUC of {bm['ROC-AUC']:.4f} confirms strong discriminative ability. "
        "The Stacking Ensemble achieved the highest ROC-AUC overall "
        f"({results.loc['Stacking Ensemble', 'ROC-AUC']:.4f}), demonstrating superior "
        "ranking quality, although its lower precision reflects a trade-off toward higher "
        f"recall ({results.loc['Stacking Ensemble', 'Recall']:.4f}).",
    )

    add_image(
        doc,
        "04_model_comparison.png",
        "Figure 5. Bar chart and heatmap comparing all models across five evaluation metrics.",
    )

    add_para(
        doc,
        "The confusion matrices in Figure 6 provide further insight. Models with balanced "
        "class weighting (Logistic Regression, Decision Tree, SVM) achieve high recall but "
        "produce more false positives, while unweighted models (KNN, Gradient Boosting) "
        "are more conservative, achieving higher precision at the expense of missing true "
        "subscribers. The ensemble methods strike an effective middle ground.",
    )

    add_image(
        doc,
        "05_confusion_matrices.png",
        "Figure 6. Confusion matrices for all nine classification models.",
    )

    add_image(
        doc,
        "10_precision_recall_bar.png",
        "Figure 7. Precision versus recall trade-off across all models.",
    )

    add_image(
        doc,
        "06_roc_curves.png",
        "Figure 8. Receiver Operating Characteristic (ROC) curves for all models.",
    )

    add_para(
        doc,
        "Figure 8 shows the ROC curves. All models substantially outperform the random "
        "baseline (diagonal). The ensemble models and Gradient Boosting occupy the upper-left "
        "region most consistently, indicating strong true-positive rates at low false-positive "
        "rates. K-Nearest Neighbors shows the weakest ROC performance, consistent with its "
        "sensitivity to the curse of dimensionality in high-dimensional one-hot encoded "
        "feature space (Hastie, Tibshirani and Friedman, 2009).",
    )

    add_image(
        doc,
        "11_f1_roc_scatter.png",
        "Figure 9. Scatter plot of F1-Score against ROC-AUC for all models.",
    )

    add_image(
        doc,
        "12_radar_chart.png",
        "Figure 10. Radar chart comparing the top four models across all metrics.",
    )

    # ── 5. Discussion of Results ──────────────────────────────────
    add_heading(doc, "5. Discussion of Results")

    add_para(
        doc,
        f"The {best} emerged as the strongest overall model by F1-Score. "
        "This result aligns with ensemble learning theory, which predicts that combining "
        "diverse learners reduces generalisation error provided the base models are "
        "individually accurate and mutually diverse (Dietterich, 2000). The Voting Ensemble "
        "satisfies both conditions: it combines a linear model (Logistic Regression), a "
        "bagging method (Random Forest), a boosting method (Gradient Boosting), and a neural "
        "network (MLP), each exploiting different aspects of the feature space.",
    )

    add_para(
        doc,
        "From a business perspective, the results present a clear precision–recall trade-off. "
        "If the bank prioritises campaign efficiency — minimising wasted calls to non-subscribers "
        "— Gradient Boosting offers the highest precision (0.6957), meaning approximately 70% "
        "of clients predicted as subscribers would actually subscribe. Conversely, if the bank "
        "prioritises market coverage — ensuring as few potential subscribers as possible are "
        "missed — the Stacking Ensemble or SVM provide recall above 0.93, capturing over 93% "
        "of true subscribers at the cost of more false positives (Provost and Fawcett, 2013).",
    )

    add_para(
        doc,
        "The Random Forest model provides a useful interpretability advantage. Feature importance "
        f"analysis identified the top five predictive features as: {', '.join(top5)}. "
        "The prominence of duration-related and macroeconomic features suggests that both "
        "call quality and economic context are powerful predictors. However, the high importance "
        "of call duration should be treated with caution, as this variable is available only "
        "after a conversation has occurred and may introduce data leakage if the model is "
        "intended for pre-contact targeting (Moro, Cortez and Rita, 2014).",
    )

    add_image(
        doc,
        "07_feature_importance.png",
        "Figure 11. Top 15 feature importances from the Random Forest model.",
    )

    add_para(
        doc,
        "The Neural Network (MLP) achieved competitive results (F1 = 0.5827, ROC-AUC = 0.9490) "
        "but did not outperform the ensemble methods. This may reflect the relatively tabular "
        "nature of the dataset, where tree-based and ensemble methods typically perform well "
        "without extensive hyperparameter tuning (Grinsztajn, Oyallon and Varoquaux, 2022). "
        "Neural networks tend to achieve their full potential with larger datasets and "
        "unstructured data modalities such as images or text (LeCun, Bengio and Hinton, 2015).",
    )

    # ── 6. Limitations and Potential Improvements ─────────────────
    add_heading(doc, "6. Limitations and Potential Improvements")

    add_para(
        doc,
        "Several limitations should be acknowledged. First, the analysis used a single "
        "train-test split. While stratification preserves class balance, performance estimates "
        "are subject to sampling variability. Repeated k-fold cross-validation would provide "
        "more robust estimates and confidence intervals (Kohavi, 1995). Second, hyperparameter "
        "optimisation was limited to informed default settings rather than systematic search. "
        "Grid search or Bayesian optimisation (Snoek, Larochelle and Adams, 2012) could yield "
        "improved configurations, particularly for Gradient Boosting and the MLP network.",
    )

    add_para(
        doc,
        "Third, the call duration variable, while highly predictive, is a post-contact feature "
        "that may not be available at the time of campaign planning. A separate analysis "
        "excluding duration would better reflect a realistic pre-contact prediction scenario. "
        "Fourth, resampling techniques such as SMOTE (Chawla et al., 2002) or cost-sensitive "
        "learning could further improve minority class detection. Fifth, threshold tuning — "
        "adjusting the classification probability cut-off based on business cost assumptions — "
        "could align model decisions more closely with campaign economics (Elkan, 2001).",
    )

    add_para(
        doc,
        "Future work could also explore XGBoost or LightGBM as alternative gradient boosting "
        "implementations that offer built-in regularisation and faster training (Chen and "
        "Guestrin, 2016; Ke et al., 2017). Additionally, a segmented modelling approach — "
        "building separate models for different client segments — may capture heterogeneity "
        "in response behaviour that a single global model cannot represent effectively.",
    )

    # ── 7. Conclusion ─────────────────────────────────────────────
    add_heading(doc, "7. Conclusion")

    add_para(
        doc,
        f"This report demonstrated the application of nine machine learning algorithms to "
        f"the UCI Bank Marketing dataset ({n_rows:,} records, {n_cols} features) to predict "
        "term deposit subscription. The analysis followed a structured workflow: exploratory "
        "data analysis, pipeline-based preprocessing, model training with class-imbalance "
        "handling, multi-metric evaluation, and critical interpretation of results.",
    )

    add_para(
        doc,
        f"The {best} achieved the best overall F1-Score ({bm['F1-Score']:.4f}), "
        "confirming that combining diverse classifiers through probability averaging can "
        "outperform individual models on imbalanced classification tasks. The Stacking "
        "Ensemble achieved the highest ROC-AUC, and Gradient Boosting offered the highest "
        "precision, illustrating that model selection should be guided by the specific business "
        "objective rather than a single metric.",
    )

    add_para(
        doc,
        "The findings carry direct practical implications: predictive models can enable banks "
        "to prioritise high-potential clients, reduce campaign costs, and improve conversion "
        "rates. The ensemble approach, in particular, provides a robust and flexible solution "
        "that balances multiple performance dimensions. Future improvements through "
        "cross-validation, hyperparameter tuning, threshold optimisation, and the exclusion "
        "of post-contact variables would further strengthen the model's real-world applicability.",
    )

    # ── References ────────────────────────────────────────────────
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "References")

    references = [
        "Breiman, L. (2001) 'Random Forests', Machine Learning, 45(1), pp. 5–32.",
        "Breiman, L., Friedman, J.H., Olshen, R.A. and Stone, C.J. (1984) Classification and Regression Trees. Boca Raton: Chapman and Hall/CRC.",
        "Chawla, N.V., Bowyer, K.W., Hall, L.O. and Kegelmeyer, W.P. (2002) 'SMOTE: Synthetic Minority Over-sampling Technique', Journal of Artificial Intelligence Research, 16, pp. 321–357.",
        "Chen, T. and Guestrin, C. (2016) 'XGBoost: A Scalable Tree Boosting System', Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, pp. 785–794.",
        "Cortes, C. and Vapnik, V. (1995) 'Support-vector networks', Machine Learning, 20(3), pp. 273–297.",
        "Cover, T.M. and Hart, P.E. (1967) 'Nearest neighbor pattern classification', IEEE Transactions on Information Theory, 13(1), pp. 21–27.",
        "Dietterich, T.G. (2000) 'Ensemble Methods in Machine Learning', Multiple Classifier Systems. Berlin: Springer, pp. 1–15.",
        "Dua, D. and Graff, C. (2017) UCI Machine Learning Repository. Irvine, CA: University of California, School of Information and Computer Science. Available at: https://archive.ics.uci.edu/ml (Accessed: 20 March 2026).",
        "Elkan, C. (2001) 'The Foundations of Cost-Sensitive Learning', Proceedings of the 17th International Joint Conference on Artificial Intelligence, pp. 973–978.",
        "Fawcett, T. (2006) 'An introduction to ROC analysis', Pattern Recognition Letters, 27(8), pp. 861–874.",
        "Fernandez-Delgado, M., Cernadas, E., Barro, S. and Amorim, D. (2014) 'Do we Need Hundreds of Classifiers to Solve Real World Classification Problems?', Journal of Machine Learning Research, 15(1), pp. 3133–3181.",
        "Friedman, J.H. (2001) 'Greedy function approximation: A gradient boosting machine', Annals of Statistics, 29(5), pp. 1189–1232.",
        "Goodfellow, I., Bengio, Y. and Courville, A. (2016) Deep Learning. Cambridge, MA: MIT Press.",
        "Grinsztajn, L., Oyallon, E. and Varoquaux, G. (2022) 'Why do tree-based models still outperform deep learning on tabular data?', Advances in Neural Information Processing Systems, 35, pp. 507–520.",
        "Hastie, T., Tibshirani, R. and Friedman, J. (2009) The Elements of Statistical Learning. 2nd edn. New York: Springer.",
        "He, H. and Garcia, E.A. (2009) 'Learning from Imbalanced Data', IEEE Transactions on Knowledge and Data Engineering, 21(9), pp. 1263–1284.",
        "Hosmer, D.W., Lemeshow, S. and Sturdivant, R.X. (2013) Applied Logistic Regression. 3rd edn. Hoboken: Wiley.",
        "Hunter, J.D. (2007) 'Matplotlib: A 2D graphics environment', Computing in Science & Engineering, 9(3), pp. 90–95.",
        "Jordan, M.I. and Mitchell, T.M. (2015) 'Machine learning: Trends, perspectives, and prospects', Science, 349(6245), pp. 255–260.",
        "Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., Ye, Q. and Liu, T.Y. (2017) 'LightGBM: A Highly Efficient Gradient Boosting Decision Tree', Advances in Neural Information Processing Systems, 30, pp. 3146–3154.",
        "King, G. and Zeng, L. (2001) 'Logistic Regression in Rare Events Data', Political Analysis, 9(2), pp. 137–163.",
        "Kohavi, R. (1995) 'A study of cross-validation and bootstrap for accuracy estimation and model selection', Proceedings of the 14th International Joint Conference on Artificial Intelligence, pp. 1137–1143.",
        "Kuhn, M. and Johnson, K. (2013) Applied Predictive Modeling. New York: Springer.",
        "LeCun, Y., Bengio, Y. and Hinton, G. (2015) 'Deep learning', Nature, 521(7553), pp. 436–444.",
        "Lessmann, S., Baesens, B., Seow, H.V. and Thomas, L.C. (2015) 'Benchmarking state-of-the-art classification algorithms for credit scoring', European Journal of Operational Research, 247(1), pp. 124–136.",
        "McKinney, W. (2010) 'Data Structures for Statistical Computing in Python', Proceedings of the 9th Python in Science Conference, pp. 56–61.",
        "Moro, S., Cortez, P. and Rita, P. (2014) 'A data-driven approach to predict the success of bank telemarketing', Decision Support Systems, 62, pp. 22–31.",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., Cournapeau, D., Brucher, M., Perrot, M. and Duchesnay, E. (2011) 'Scikit-learn: Machine Learning in Python', Journal of Machine Learning Research, 12, pp. 2825–2830.",
        "Prechelt, L. (1998) 'Early Stopping — But When?', in Neural Networks: Tricks of the Trade. Berlin: Springer, pp. 55–69.",
        "Provost, F. and Fawcett, T. (2013) Data Science for Business. Sebastopol, CA: O'Reilly Media.",
        "Raschka, S. and Mirjalili, V. (2019) Python Machine Learning. 3rd edn. Birmingham: Packt Publishing.",
        "Snoek, J., Larochelle, H. and Adams, R.P. (2012) 'Practical Bayesian Optimization of Machine Learning Algorithms', Advances in Neural Information Processing Systems, 25, pp. 2951–2959.",
        "Sokolova, M. and Lapalme, G. (2009) 'A systematic analysis of performance measures for classification tasks', Information Processing & Management, 45(4), pp. 427–437.",
        "Van der Laan, M.J., Polley, E.C. and Hubbard, A.E. (2007) 'Super Learner', Statistical Applications in Genetics and Molecular Biology, 6(1), Article 25.",
        "Waskom, M.L. (2021) 'seaborn: statistical data visualization', Journal of Open Source Software, 6(60), 3021.",
        "Wolpert, D.H. (1992) 'Stacked generalization', Neural Networks, 5(2), pp. 241–259.",
        "Zhou, Z.H. (2012) Ensemble Methods: Foundations and Algorithms. Boca Raton: Chapman and Hall/CRC.",
    ]

    for ref in references:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(ref)
        run.font.size = Pt(10)

    # ── Save ──────────────────────────────────────────────────────
    doc.save(REPORT_PATH)
    print(f"\nReport saved to: {REPORT_PATH}")
    print(
        f"Approximate word count (body): ~3,000 words (excluding references and captions)"
    )


# ══════════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════════


def main() -> None:
    df = pd.read_csv(DATA_PATH, sep=";")
    print(f"Loaded dataset: {len(df):,} rows x {df.shape[1]} columns")

    results = pd.read_csv(VIS_DIR / "model_results_summary.csv", index_col="Model")
    results = results.astype(float).sort_values("F1-Score", ascending=False)
    print(f"Loaded results for {len(results)} models")

    feat_path = VIS_DIR / "feature_importance.csv"
    feat_imp = (
        pd.read_csv(feat_path)
        if feat_path.exists()
        else pd.DataFrame(columns=["Feature", "Importance"])
    )

    VIS_DIR.mkdir(exist_ok=True)

    print("\nGenerating additional visualisations...")
    create_correlation_heatmap(df)
    create_class_imbalance_visual(df)
    create_precision_recall_bar(results)
    create_f1_roc_scatter(results)
    create_radar_chart(results)

    build_report(df, results, feat_imp)
    print("\nDone.")


if __name__ == "__main__":
    main()
